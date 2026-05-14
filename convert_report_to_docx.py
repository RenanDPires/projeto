#!/usr/bin/env python3
"""
Script para converter relatório Markdown para .docx com formatação acadêmica.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_background(cell, fill):
    """Define cor de fundo de célula."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_report_docx():
    """Cria documento Word com relatório formatado."""
    
    doc = Document()
    
    # Configurar espaçamento de parágrafo padrão
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ===== CAPA / CABEÇALHO =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('RELATÓRIO DE IMPLEMENTAÇÃO')
    run.font.size = Pt(16)
    run.font.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Eletromag Lab - Plataforma Interativa de Análise Eletromagnética')
    run.font.size = Pt(12)
    run.font.bold = True
    
    doc.add_paragraph()  # espaço
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run('Data: 14 de Maio de 2026')
    run.font.size = Pt(10)
    run.italic = True
    
    doc.add_page_break()
    
    # ===== ÍNDICE =====
    doc.add_heading('ÍNDICE', level=1)
    toc_items = [
        '1. INTRODUÇÃO',
        '2. ABORDAGEM E METODOLOGIA',
        '3. DETALHES POR QUESTÃO',
        '4. AMBIENTE IMPLEMENTADO',
        '5. VALIDAÇÃO E TESTES',
        '6. COMPILAÇÃO E DISTRIBUIÇÃO',
        '7. DOCUMENTAÇÃO E REFERÊNCIAS',
        '8. LIÇÕES APRENDIDAS',
        '9. ANEXO: RECORTES DO CÓDIGO',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 1. INTRODUÇÃO =====
    doc.add_heading('1. INTRODUÇÃO', level=1)
    doc.add_paragraph(
        'O projeto Eletromag Lab é uma aplicação web interativa desenvolvida em Python '
        'utilizando o framework Streamlit, com o objetivo de fornecer ferramentas '
        'computacionais para análise e solução de problemas em eletromagnetismo aplicado, '
        'especificamente focado em transformadores de potência e condutores em campos '
        'eletromagnéticos.'
    )
    doc.add_paragraph(
        'A plataforma implementa as cinco questões de uma avaliação acadêmica '
        '(EEL7216 08202 - Tópicos Especiais em Eletromagnética):'
    )
    
    intro_items = [
        'Análise de perdas em tanques (Q1)',
        'Soluções analíticas em difusão de campo magnético (Q2)',
        'Cálculo de campo magnético via Biot-Savart (Q3)',
        'Resistência AC em condutores retangulares (Q4)',
        'Comparação de métodos analíticos (Q5)',
    ]
    for item in intro_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # ===== 2. ABORDAGEM =====
    doc.add_heading('2. ABORDAGEM E METODOLOGIA DO PROJETO', level=1)
    
    doc.add_heading('2.1 Stack Tecnológico', level=2)
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Cabeçalho
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Componente'
    hdr_cells[1].text = 'Versão'
    hdr_cells[2].text = 'Propósito'
    for cell in hdr_cells:
        set_cell_background(cell, 'D3D3D3')
    
    # Dados
    tech_data = [
        ('Python', '3.12.10', 'Linguagem base'),
        ('Streamlit', '1.55.0', 'Interface web interativa'),
        ('NumPy', '1.26.0+', 'Computação numérica vetorizada'),
        ('SciPy', '1.11.0+', 'Cálculos científicos'),
        ('Plotly', '5.17.0+', 'Visualização 2D/3D'),
        ('PyInstaller', '6.20.0', 'Compilação para executável'),
        ('Pandas', '-', 'Análise de dados e CSV'),
    ]
    
    for idx, (comp, vers, prop) in enumerate(tech_data, start=1):
        row_cells = table.rows[idx].cells
        row_cells[0].text = comp
        row_cells[1].text = vers
        row_cells[2].text = prop
    
    doc.add_heading('2.2 Princípios de Design', level=2)
    doc.add_paragraph(
        'Separação Frontend-Backend: O código de cálculo eletromagnético (app/core/) '
        'é totalmente independente da renderização Streamlit (app/tabs/).'
    )
    doc.add_paragraph(
        'Modularização por Domínio: Cada área física é uma classe separada, testável isoladamente.'
    )
    doc.add_paragraph(
        'Validação via Pydantic: Todos os inputs e outputs validados por schemas Pydantic.'
    )
    
    # ===== 3. DETALHES POR QUESTÃO =====
    doc.add_heading('3. DETALHES DE IMPLEMENTAÇÃO POR QUESTÃO', level=1)
    
    doc.add_heading('3.1 Questão 1: Análise de Perdas em Tanques', level=2)
    doc.add_paragraph(
        'Cálculo volumétrico de perdas de potência em tanques cilíndricos utilizando '
        'integração Gauss-Legendre em 3D.'
    )
    doc.add_paragraph('Status: ✅ Completamente desenvolvida')
    
    doc.add_heading('3.2 Questão 2: Soluções Analíticas em Difusão Magnética', level=2)
    doc.add_paragraph(
        'Solução da equação de difusão de campo magnético em placas condutoras.'
    )
    warning = doc.add_paragraph()
    run = warning.add_run('⚠️ NOTA IMPORTANTE: ')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    warning.add_run(
        'A Q2 NÃO foi desenvolvida de forma completamente independente. '
        'O código utiliza principalmente formulações de referência do artigo Kaymak et al. (2016) '
        'e adaptações de soluções já estabelecidas na literatura.'
    )
    doc.add_paragraph('Status: ✅ Implementada (com base em literatura consolidada)')
    
    doc.add_heading('3.3 Questão 3: Campo Magnético via Biot-Savart', level=2)
    doc.add_paragraph(
        'Cálculo vetorial de campo magnético gerado por geometrias de condutor '
        '(espiras, solenoides, helicoidais).'
    )
    doc.add_paragraph('Características: Suporte a múltiplas geometrias, quadratura adaptativa, visualização 3D.')
    doc.add_paragraph('Status: ✅ Completamente desenvolvida')
    
    doc.add_heading('3.4 Questão 4: Resistência AC em Condutores Retangulares', level=2)
    doc.add_paragraph(
        'Cálculo de perdas em condutores retangulares sob diferentes condições de contorno (Dowell method).'
    )
    doc.add_paragraph('Três variantes implementadas:')
    doc.add_paragraph('(a) Simétrico: Campo em ambas superfícies', style='List Bullet')
    doc.add_paragraph('(b) Baseline: Semi-infinito (referência)', style='List Bullet')
    doc.add_paragraph('(c) Finito: Campo confinado entre superfícies', style='List Bullet')
    
    warning2 = doc.add_paragraph()
    run = warning2.add_run('⚠️ NOTA IMPORTANTE: ')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    warning2.add_run(
        'A Q4 NÃO foi desenvolvida como solução original independente. '
        'A implementação baseia-se em formulações fechadas do método Dowell (1966) '
        'publicadas em Del Vecchio (2010) e Kulkarni (2013).'
    )
    doc.add_paragraph('Status: ✅ Implementada (equações de literatura validadas)')
    
    doc.add_heading('3.5 Questão 5: Comparação de Métodos Analíticos', level=2)
    doc.add_paragraph(
        'Comparação sistemática de perdas em diferentes geometrias de condutores '
        '(circular, retangular variantes, sheet).'
    )
    doc.add_paragraph(
        'Recentes Melhorias (Maio 2026): Renovação da seção de Interpretação Física, '
        'refatoração dos gráficos para mostrar variação de geometria, novos gráficos de varredura.'
    )
    doc.add_paragraph('Status: ✅ Completamente desenvolvida e atualizada')
    
    # ===== 4. AMBIENTE =====
    doc.add_heading('4. AMBIENTE IMPLEMENTADO - ESTRUTURA DE INTERFACE', level=1)
    
    doc.add_heading('4.1 Estrutura de Abas Principal', level=2)
    doc.add_paragraph(
        'A aplicação Streamlit organiza-se em uma aba de apresentação seguida de '
        'cinco abas principais (Q1-Q5). Cada aba contém:'
    )
    doc.add_paragraph('Aba Teoria: Equações LaTeX, referências, conceitos fundamentais', style='List Bullet')
    doc.add_paragraph('Aba Cálculo: Inputs interativos, botão de execução', style='List Bullet')
    doc.add_paragraph('Aba Resultados: Tabelas, gráficos, análise física', style='List Bullet')
    
    doc.add_heading('4.2 Exemplo: Questão 5', level=2)
    doc.add_paragraph(
        'A Q5 permite comparar perdas em diferentes geometrias. O usuário fornece:'
    )
    doc.add_paragraph('Dimensão característica do condutor', style='List Bullet')
    doc.add_paragraph('Condutividade do material', style='List Bullet')
    doc.add_paragraph('Frequência de operação', style='List Bullet')
    doc.add_paragraph('Campo magnético superficial', style='List Bullet')
    
    doc.add_paragraph('O sistema exibe:')
    doc.add_paragraph('Tabela comparativa de perdas [W/m²] para cada geometria', style='List Bullet')
    doc.add_paragraph('Interpretação física dos resultados', style='List Bullet')
    doc.add_paragraph('Gráficos interativos: Perdas vs Frequência e Perdas vs Dimensão', style='List Bullet')
    doc.add_paragraph('Exportação em CSV dos dados', style='List Bullet')
    
    # ===== 5. VALIDAÇÃO =====
    doc.add_heading('5. VALIDAÇÃO E TESTES', level=1)
    
    doc.add_paragraph(
        'Suite de testes desenvolvida com pytest, cobrindo 22 testes específicos para Q5 '
        'e totalizando 104 testes globais do projeto.'
    )
    doc.add_paragraph('Validação inclui:')
    doc.add_paragraph('Casos limites com soluções analíticas conhecidas', style='List Bullet')
    doc.add_paragraph('Análise dimensional de todas as equações', style='List Bullet')
    doc.add_paragraph('Comparação com valores publicados em literatura', style='List Bullet')
    
    doc.add_paragraph()
    result = doc.add_paragraph('Resultados: ')
    run = result.add_run('✅ 104/104 testes passando')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    
    # ===== 6. COMPILAÇÃO =====
    doc.add_heading('6. COMPILAÇÃO E DISTRIBUIÇÃO', level=1)
    
    doc.add_heading('6.1 Executável Autônomo', level=2)
    doc.add_paragraph('O projeto foi compilado para Windows 64-bit usando PyInstaller.')
    
    table2 = doc.add_table(rows=8, cols=2)
    table2.style = 'Light Grid Accent 1'
    
    spec_data = [
        ('Arquivo', 'dist/EEL7216_Lab.exe'),
        ('Tamanho', '128.8 MB'),
        ('Compatibilidade', 'Windows 10/11 (64-bit)'),
        ('Dependência Python', 'Nenhuma (incluso)'),
        ('Primeira execução', '~20-30 segundos'),
        ('Próximas execuções', 'Rápidas (cache)'),
        ('Tempo de compilação', '~3-4 minutos'),
    ]
    
    hdr = table2.rows[0].cells
    hdr[0].text = 'Propriedade'
    hdr[1].text = 'Valor'
    for cell in hdr:
        set_cell_background(cell, 'D3D3D3')
    
    for idx, (prop, val) in enumerate(spec_data, start=1):
        row = table2.rows[idx].cells
        row[0].text = prop
        row[1].text = val
    
    # ===== 7. REFERÊNCIAS =====
    doc.add_heading('7. DOCUMENTAÇÃO E REFERÊNCIAS', level=1)
    
    doc.add_heading('7.1 Referências Principais', level=2)
    
    ref1 = doc.add_paragraph()
    ref1.add_run('Kaymak et al. (2016): ').bold = True
    ref1.add_run(
        'Comparison of Analytical Methods for Calculating AC Resistance. '
        'IEEE Transactions on Industry Applications, Vol. 52, No. 5.'
    )
    
    ref2 = doc.add_paragraph()
    ref2.add_run('Del Vecchio et al. (2010): ').bold = True
    ref2.add_run(
        'Transformer Design Principles (2nd ed.). '
        'Seção 15: AC Resistance em Condutores.'
    )
    
    ref3 = doc.add_paragraph()
    ref3.add_run('Kulkarni & Khaparde (2013): ').bold = True
    ref3.add_run(
        'Transformer Engineering: Design and Practice. '
        'Seção 4: Resistência e Perdas em Transformadores.'
    )
    
    # ===== 8. CONCLUSÃO =====
    doc.add_heading('8. CONCLUSÃO FINAL', level=1)
    
    doc.add_paragraph(
        'O projeto Eletromag Lab apresenta uma implementação robusta e educacional '
        'de conceitos eletromagnéticos avançados, com interface interativa e validação rigorosa.'
    )
    
    doc.add_heading('Status das Questões:', level=2)
    doc.add_paragraph('✅ Q1, Q3, Q5: Desenvolvidas de forma independente com métodos próprios')
    doc.add_paragraph(
        '⚠️ Q2, Q4: Implementadas baseando-se em soluções publicadas de literatura consolidada, '
        'sem derivação original de método alternativo'
    )
    
    doc.add_paragraph()
    final = doc.add_paragraph('Documento Gerado: 14 de Maio de 2026 | Versão: 128.8 MB')
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.runs[0].italic = True
    
    # Salvar documento
    output_path = r'C:\Users\renan\OneDrive\Área de Trabalho\Renan\UFSC\Mauricio\RELATORIO_ELETROMAG_LAB.docx'
    doc.save(output_path)
    print(f'✓ Documento criado: {output_path}')
    return output_path


if __name__ == '__main__':
    create_report_docx()
